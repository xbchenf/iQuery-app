## 导入依赖

import os


import tempfile
import shutil
from docx import Document

base_path = "D:\\yuanma\\iquery\\iquery云盘"
class CloudFile():
    """
    操作云文档
    """

    def __init__(self,
                 project_name,
                 part_name,
                 doc_content=None):

        # 项目名称，即项目文件夹名称
        self.project_name = project_name
        # 项目某部分名称，即项目文件名称
        self.part_name = part_name

        folder_path = create_or_get_folder(folder_name=project_name)

        # 创建时获取当前项目中其他文件名称列表
        self.doc_list = list_files_in_folder(folder_name=project_name)

        file_path = create_or_get_doc(folder_name=self.project_name,
                                      doc_name=self.part_name)

        # 项目文件具体内容，相当于多轮对话内容
        self.doc_content = doc_content
        # 若初始content不为空，则将其追加入文档内
        if doc_content != None:
            append_content_in_doc(folder_name=project_name,
                                  doc_name=part_name,
                                  dict_list=doc_content)

    def get_doc_content(self):
        """
        根据项目某文件的文件ID，获取对应的文件内容
        """
        self.doc_content = get_file_content(folder_name=self.project_name, doc_name=self.part_name)

        return self.doc_content

    def append_doc_content(self, content):
        """
        根据项目某文件的文件ID，追加文件内容
        """
        append_content_in_doc(folder_name=self.project_name,
                              doc_name=self.part_name,
                              dict_list=content)

    def clear_content(self):
        """
        清空某文件内的全部内容
        """
        clear_content_in_doc(folder_name=self.project_name, doc_name=self.part_name)

    def delete_all_files(self):
        """
        删除当前项目文件夹内的全部文件
        """
        delete_all_files_in_folder(folder_name=self.project_name)

    def update_doc_list(self):
        """
        更新当前项目文件夹内的全部文件名称
        """
        self.doc_list = list_files_in_folder(folder_name=self.project_name)

    def rename_doc(self, new_name):
        """
        修改当前文件名称
        """
        self.part_name = rename_doc(folder_name=self.project_name, doc_name=self.part_name,
                                             new_name=new_name)

def create_or_get_folder(folder_name):
    """
    根据项目创建云盘目录
    """
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    full_path = os.path.join(base_path, folder_name)
    # 如果目录不存在，则创建它
    if not os.path.exists(full_path):
        os.makedirs(full_path)
        print(f"目录 {folder_name} 创建成功")
    else:
        print(f"目录 {folder_name} 已存在")

def create_or_get_doc(folder_name, doc_name):
    """
    创建或获取文件路径
    """
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    full_path_folder = os.path.join(base_path, folder_name)
    file_path_doc = os.path.join(base_path + "/" + folder_name, f'{doc_name}.doc')

    # 检查目录是否存在，如果不存在则创建
    if not os.path.exists(full_path_folder):
        os.makedirs(full_path_folder)

    # 检查文件是否存在
    if os.path.exists(file_path_doc):
        # 文件存在，打开并追加内容
        document = Document(file_path_doc)
    else:
        # 文件不存在，创建一个新的文档对象
        document = Document()
    # 保存文档
    document.save(file_path_doc)

    return file_path_doc

def append_content_in_doc(folder_name, doc_name, qa_string):
    """"
    往文件里追加内容
    @param folder_name=目录名，doc_name=文件名，qa_string=追加的内容
    """
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    ## 目录地址
    full_path_folder = base_path + "/" + folder_name
    ## 文件地址
    full_path_doc = os.path.join(full_path_folder, doc_name) + ".doc"

    # 检查目录是否存在，如果不存在则创建
    if not os.path.exists(full_path_folder):
        os.makedirs(full_path_folder)

    # 检查文件是否存在
    if os.path.exists(full_path_doc):
        # 文件存在，打开并追加内容
        document = Document(full_path_doc)
    else:
        # 文件不存在，创建一个新的文档对象
        document = Document()
    # 追加内容
    document.add_paragraph(qa_string)
    # 保存文档
    document.save(full_path_doc)
    print(f"内容已追加到 {doc_name}")

def get_file_content(folder_name, doc_name):
    """
    实现根据项目名和文件名获取文件内容的方法
    @param project_name:项目名，file_name：文件名
    @return 文件内容
    """
    # 构建文件的完整路径
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    file_path = os.path.join(folder_name, doc_name)
    full_path = os.path.join(base_path, file_path) + ".doc"

    # 确保文件存在
    if not os.path.exists(full_path):
        return "文件不存在"

    try:
        # 加载文档
        doc = Document(full_path)
        content = []

        # 遍历文档中的每个段落，并收集文本
        for para in doc.paragraphs:
            content.append(para.text)

        # 将所有段落文本合并成一个字符串返回
        return '\n'.join(content)
    except Exception as e:
        return f"读取文件时发生错误: {e}"

def clear_content_in_doc(folder_name, doc_name):
    # 打开文档
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    file_path = os.path.join(base_path + "/" + folder_name, f'{doc_name}.doc')
    doc = Document(file_path)

    # 遍历每一个段落，设置其文本为空字符串
    for p in doc.paragraphs:
        for run in p.runs:
            run.text = ''

    # 保存修改后的文档
    doc.save(file_path)
    print("文档内容清除完毕")

def list_files_in_folder(folder_name):
    """
    列举当前文件夹的全部文件
    """
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    full_path = os.path.join(base_path, folder_name)
    file_names = [f for f in os.listdir(full_path) if os.path.isfile(os.path.join(full_path, f))]

    return file_names

def rename_doc(folder_name, doc_name, new_name):
    """
    修改指定的文档名称
    """
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    file_path = os.path.join(base_path + "/" + folder_name, f'{doc_name}.doc')
    new_file_path = os.path.join(base_path + "/" + folder_name, f'{new_name}.doc')
    # 重命名文件
    os.rename(file_path, new_file_path)

    return new_name

def delete_all_files_in_folder(folder_name):
    """
    删除某文件夹内全部文件
    """
    # 定义要删除的目录路径
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    full_path = os.path.join(base_path, folder_name)
    # 遍历整个目录
    for filename in os.listdir(full_path):
        # 构造文件或者文件夹的绝对路径
        file_path = os.path.join(full_path, filename)
        try:
            # 如果是文件，则删除文件
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            # 如果是文件夹，则删除文件夹
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
                print("文件已清除完毕")
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))

def append_img_in_doc(folder_name, doc_name, fig):
    """"
    往文件里追加图片
    @param folder_name=目录名，doc_name=文件名，img=图片对象，数据类型为matplotlib.figure.Figure对象
    """
    #base_path = "/root/autodl-tmp/iquery项目/iquery云盘"
    ## 目录地址
    full_path_folder = base_path + "/" + folder_name
    ## 文件地址
    full_path_doc = os.path.join(full_path_folder, doc_name) + ".doc"

    # 检查目录是否存在，如果不存在则创建
    if not os.path.exists(full_path_folder):
        os.makedirs(full_path_folder)

    # 检查文件是否存在
    if os.path.exists(full_path_doc):
        print(full_path_doc)
         # 文件存在，打开并追加内容
        document = Document(full_path_doc)
    else:
        # 文件不存在，创建一个新的文档对象
        document = Document()

    # 追加图片
    # 将matplotlib的Figure对象保存为临时图片文件
    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
        fig.savefig(tmpfile.name, format='png')
        # 将图片插入到.docx文档中
        document.add_picture(tmpfile.name)

    # 保存文档
    document.save(full_path_doc)
    print(f"图片已追加到 {doc_name}")

